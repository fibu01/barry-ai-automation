"""
Barry University - AI Tip of the Week generator.
Run: python3 create_totw.py
Requires env vars: TENANT_ID, CLIENT_ID, CLIENT_SECRET, SHAREPOINT_DRIVE_ID
"""

import os
import re
import time
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SHAREPOINT_FOLDER = "AI Governance/Tip of the Week"


# ---------------------------------------------------------------------------
# Auth
# ---------------------------------------------------------------------------

def authenticate():
    tenant_id = os.environ["TENANT_ID"]
    client_id = os.environ["CLIENT_ID"]
    client_secret = os.environ["CLIENT_SECRET"]

    resp = requests.post(
        f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token",
        data={
            "grant_type": "client_credentials",
            "client_id": client_id,
            "client_secret": client_secret,
            "scope": "https://graph.microsoft.com/.default",
        },
    )
    if resp.status_code != 200:
        raise Exception(f"AUTH FAILED: {resp.text}")
    print("AUTH SUCCESS")
    return resp.json()["access_token"]


# ---------------------------------------------------------------------------
# SharePoint helpers
# ---------------------------------------------------------------------------

def graph_get(token, url, stream=False, retries=4):
    headers = {"Authorization": f"Bearer {token}"}
    wait = 2
    for attempt in range(1, retries + 1):
        r = requests.get(url, headers=headers, stream=stream)
        if r.status_code == 200:
            return r
        if attempt < retries:
            print(f"  GET attempt {attempt} returned {r.status_code}, retrying in {wait}s...")
            time.sleep(wait)
            wait *= 2
    return r  # return last failed response for caller to inspect


def download_banner(token, drive_id, tip_num):
    # Try both zero-padded (TOTW_05_banner.png) and non-padded (TOTW_5_banner.png)
    candidates = [
        f"TOTW_{tip_num:02d}_banner.png",
        f"TOTW_{tip_num}_banner.png",
    ]
    dest = f"/tmp/{candidates[0]}"  # always save under zero-padded name

    wait = 2
    for attempt in range(1, 5):
        for filename in candidates:
            url = (
                f"https://graph.microsoft.com/v1.0/drives/{drive_id}"
                f"/root:/{SHAREPOINT_FOLDER}/{filename}:/content"
            )
            headers = {"Authorization": f"Bearer {token}"}
            r = requests.get(url, headers=headers, stream=True)
            if r.status_code == 200:
                with open(dest, "wb") as f:
                    for chunk in r.iter_content(chunk_size=8192):
                        f.write(chunk)
                print(f"BANNER DOWNLOADED: {filename} ({os.path.getsize(dest)} bytes)")
                return dest
            print(f"  Attempt {attempt}: {filename} returned {r.status_code}")

        if attempt < 4:
            print(f"  Banner not available, retrying in {wait}s...")
            time.sleep(wait)
            wait *= 2

    raise Exception(
        f"BANNER DOWNLOAD FAILED: Could not retrieve banner for TIP #{tip_num} "
        f"after 4 attempts. Tried: {candidates}. Document not created."
    )


def upload_file(token, drive_id, local_path, remote_filename):
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/octet-stream",
    }
    upload_url = (
        f"https://graph.microsoft.com/v1.0/drives/{drive_id}"
        f"/root:/{SHAREPOINT_FOLDER}/{remote_filename}:/content"
    )
    with open(local_path, "rb") as f:
        data = f.read()

    wait = 2
    for attempt in range(1, 5):
        resp = requests.put(upload_url, headers=headers, data=data)
        if resp.status_code in [200, 201]:
            return resp.json().get("webUrl")
        if attempt < 4:
            print(f"  Upload attempt {attempt} failed ({resp.status_code}), retrying in {wait}s...")
            time.sleep(wait)
            wait *= 2
    raise Exception(f"UPLOAD FAILED: {resp.status_code} - {resp.text}")


# ---------------------------------------------------------------------------
# CLAUDE.md parsing
# ---------------------------------------------------------------------------

def parse_claude_md(path="CLAUDE.md", force_tip=None):
    with open(path, encoding="utf-8") as f:
        content = f.read()

    tip_blocks = re.split(r"(?=^## TIP #\d+)", content, flags=re.MULTILINE)

    if force_tip is not None:
        # Find the specific tip regardless of status (used for reruns after rejection)
        matched = [(int(m.group(1)), block)
                   for block in tip_blocks
                   if (m := re.match(r"## TIP #(\d+)", block)) and int(m.group(1)) == force_tip]
        if not matched:
            raise Exception(f"TIP #{force_tip} not found in CLAUDE.md")
        tip_num, block = matched[0]
    else:
        ready_tips = []
        for block in tip_blocks:
            m = re.match(r"## TIP #(\d+)\s*\nSTATUS:\s*(\w+)", block)
            if m and m.group(2).strip() == "Ready":
                ready_tips.append((int(m.group(1)), block))
        if not ready_tips:
            raise Exception("No tips with STATUS: Ready found in CLAUDE.md")
        tip_num, block = min(ready_tips, key=lambda x: x[0])

    def field(name):
        m = re.search(rf"^{re.escape(name)}:\s*(.+?)(?=\n[A-Z]|\Z)", block, re.MULTILINE | re.DOTALL)
        if not m:
            raise Exception(f"Field '{name}' not found in TIP #{tip_num}")
        return m.group(1).strip()

    tip = {
        "num": tip_num,
        "title": field("TITLE"),
        "subtitle": field("SUBTITLE"),
        "rethinking_title": field("RETHINKING TITLE"),
        "rethinking_body": field("RETHINKING BODY"),
        "how_to_try_intro": field("HOW TO TRY IT INTRO"),
        "bullet1": field("BULLET 1"),
        "bullet2": field("BULLET 2"),
        "bullet3": field("BULLET 3"),
        "follow_up": field("FOLLOW UP LINE"),
        "challenge": field("WEEKLY CHALLENGE"),
        "get_started": field("GET STARTED"),
    }
    return tip


def mark_done(tip_num, path="CLAUDE.md"):
    with open(path, encoding="utf-8") as f:
        content = f.read()
    updated = re.sub(
        rf"(## TIP #{tip_num}\s*\nSTATUS:)\s*Ready",
        r"\1 Done",
        content,
    )
    with open(path, "w", encoding="utf-8") as f:
        f.write(updated)
    print(f"CLAUDE.md updated: TIP #{tip_num} marked Done")


# ---------------------------------------------------------------------------
# Word document builder
# ---------------------------------------------------------------------------

def sp(para, before=0, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)


BLUE = RGBColor(0x0F, 0x47, 0x61)


def build_document(tip, banner_path, out_path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    doc.add_picture(banner_path, width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(doc.paragraphs[-1], before=0, after=8)

    p = doc.add_paragraph()
    r = p.add_run(tip["title"])
    r.bold = True; r.font.size = Pt(16); r.font.name = "Aptos"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p, before=0, after=2)

    p = doc.add_paragraph()
    r = p.add_run(tip["subtitle"])
    r.font.size = Pt(11); r.font.name = "Aptos"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(p, before=0, after=10)

    p = doc.add_paragraph()
    r = p.add_run(tip["rethinking_title"])
    r.bold = True; r.font.size = Pt(12); r.font.name = "Aptos"
    r.font.color.rgb = BLUE
    sp(p, before=0, after=4)

    p = doc.add_paragraph()
    r = p.add_run(tip["rethinking_body"])
    r.font.size = Pt(11); r.font.name = "Aptos"
    sp(p, before=0, after=10)

    p = doc.add_paragraph()
    r = p.add_run("How to Try It")
    r.bold = True; r.font.size = Pt(12); r.font.name = "Aptos"
    r.font.color.rgb = BLUE
    sp(p, before=0, after=4)

    p = doc.add_paragraph()
    r = p.add_run(tip["how_to_try_intro"])
    r.font.size = Pt(11); r.font.name = "Aptos"
    sp(p, before=0, after=4)

    for bullet in [tip["bullet1"], tip["bullet2"], tip["bullet3"]]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(bullet)
        r.font.size = Pt(11); r.font.name = "Aptos"
        sp(p, before=0, after=3)

    p = doc.add_paragraph()
    r = p.add_run(tip["follow_up"])
    r.font.size = Pt(11); r.font.name = "Aptos"
    sp(p, before=4, after=10)

    p = doc.add_paragraph()
    r = p.add_run("Weekly Challenge")
    r.bold = True; r.font.size = Pt(12); r.font.name = "Aptos"
    r.font.color.rgb = BLUE
    sp(p, before=0, after=4)

    p = doc.add_paragraph()
    r = p.add_run(tip["challenge"])
    r.font.size = Pt(11); r.font.name = "Aptos"
    sp(p, before=0, after=10)

    p = doc.add_paragraph()
    r = p.add_run("Get Started")
    r.bold = True; r.font.size = Pt(12); r.font.name = "Aptos"
    r.font.color.rgb = BLUE
    sp(p, before=0, after=4)

    p = doc.add_paragraph()
    r = p.add_run(tip["get_started"])
    r.font.size = Pt(11); r.font.name = "Aptos"
    sp(p, before=0, after=0)

    doc.save(out_path)
    print("DOCUMENT BUILT SUCCESSFULLY")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument("--tip", type=int, default=None,
                        help="Force a specific tip number (e.g. after rejection)")
    args = parser.parse_args()

    token = authenticate()
    drive_id = os.environ["SHAREPOINT_DRIVE_ID"]

    tip = parse_claude_md(force_tip=args.tip)
    tip_num = tip["num"]
    if args.tip:
        print(f"Rerunning tip: #{tip_num} - {tip['title']} (forced)")
    else:
        print(f"Current tip: #{tip_num} - {tip['title']}")

    banner_path = download_banner(token, drive_id, tip_num)

    title_slug = tip["title"].replace(" ", "_")
    doc_filename = f"TOTW_{tip_num:02d}_-_{title_slug}.docx"
    local_doc = f"/tmp/{doc_filename}"

    build_document(tip, banner_path, local_doc)

    web_url = upload_file(token, drive_id, local_doc, doc_filename)
    print(f"UPLOAD COMPLETE: {web_url}")

    if not args.tip:
        mark_done(tip_num)

    print()
    print("WORKFLOW COMPLETE")
    print(f"Tip: #{tip_num} - {tip['title']}")
    print(f"Document: {doc_filename}")
    print(f"SharePoint URL: {web_url}")
    print("Power Automate will trigger the approval workflow automatically.")


if __name__ == "__main__":
    main()
